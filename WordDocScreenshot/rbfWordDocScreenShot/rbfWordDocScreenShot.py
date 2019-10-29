from docx import Document
from docx.shared import Inches
from robot.api import logger
from robot.api.deco import keyword
import shutil
import os

class rbfWordDocScreenShot():
    
    @keyword
    def createDoc(self,title,scPath):
        document = Document()
        #self.document.add_heading(title,0)        
        for fileName in os.listdir(scPath):
            if fileName.endswith('.png'):
                pg = document.add_paragraph()
                r = pg.add_run()
                logger.info("Word file created in "+scPath+"\\"+fileName)
                r.add_picture(scPath+"\\"+fileName,width=Inches(5.8))
                logger.info("Image "+ fileName + " has been added in word file")
                #os.remove(scPath+"\\"+fileName)
        document.save(title+'.docx')

    @keyword
    def cleanScDir(self,scPath):
        for fileName in os.listdir(scPath):
            if fileName.endswith('.png'):
                os.remove(scPath+"\\"+fileName)
