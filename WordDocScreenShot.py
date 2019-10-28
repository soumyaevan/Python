from docx import Document
from docx.shared import Inches
import shutil
import os
class WordDocScreenShot():
    
    def createDoc(self,title,scPath):
        document = Document()
        #self.document.add_heading(title,0)        
        for fileName in os.listdir(scPath):
            if fileName.endswith('.png'):
                pg = document.add_paragraph()
                r = pg.add_run()
                print(scPath+"\\"+fileName)
                r.add_picture(scPath+"\\"+fileName,width=Inches(5.8))
                #os.remove(scPath+"\\"+fileName)
        document.save(title+'.docx')

    def cleanScDir(self,scPath):
        for fileName in os.listdir(scPath):
            if fileName.endswith('.png'):
                os.remove(scPath+"\\"+fileName)
