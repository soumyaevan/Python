from docx import Document
from docx.shared import Inches
import shutil
import os
class WordDoc:
    document = Document()
    def __init__(self,title,scPath):
        #self.document.add_heading(title,0)        
        for fileName in os.listdir(scPath):
            if fileName.endswith('.png'):
                pg = self.document.add_paragraph()
                r = pg.add_run()
                r.add_picture(scPath+"\\"+fileName,width=Inches(5.8))
                os.remove(scPath+"\\"+fileName)
        self.document.save(title+'.docx')

doc = WordDoc("demo","C:\\Robot1_7_Workspace\\DNB_CPP\\Reports\\Screenshots\\Login_And_Push_File")